#!/usr/bin/env python
#encoding:utf-8
#jun.yan@majorbio.com
#20170327

import argparse
import commands
import os
import re
import sys
from docx import Document
from PIL import Image

parser = argparse.ArgumentParser(
    description="iTRAQ prok report automatically generate scripts")
parser.add_argument(
    "-org", dest="org", type=str, required=True, help="'euk' or 'prok'")
parser.add_argument(
    "-m",
    dest="M1",
    type=str,
    required=True,
    help="test_department report file")
parser.add_argument(
    "-f",
    dest="file",
    type=str,
    help="please input file,normally 'Result_Files'")
parser.add_argument(
    "-up",
    dest="up",
    type=str,
    required=True,
    help="please input file,normally &default :'Result_Files'")
parser.add_argument(
    "-dump",
    dest="dump",
    type=str,
    required=True,
    help="choose you data format{dump?T:F}")
args = parser.parse_args()

if not args.file:
    args.file = 'Result_Files'

###打开各个输入文档
docx1 = Document(args.M1)
docx2 = Document(
    '/mnt/ilustre/users/jun.yan/scripts/pm_report_module/iTRAQ2.docx')
docx3 = Document(
    '/mnt/ilustre/users/jun.yan/scripts/pm_report_module/iTRAQ3-%s-%s.docx' %
    (args.org, args.dump))


###首页信息输入
def _first_page(docx):
    text1 = '\n'.join(
        [docx1.paragraphs[i].text for i in range(len(docx1.paragraphs))])
    name = re.findall(u'客户姓名：(.*)', text1)[0].strip()
    MJnum = re.findall(u'项目编号：(.*)', text1)[0].strip()
    doctime = re.findall(u'时    间：(.*)', text1)[0].strip()
    for i in range(len(docx.paragraphs)):
        if u'客户姓名' in docx.paragraphs[i].text:
            docx.paragraphs[i].runs[1].text = name
        if u'项目编号' in docx.paragraphs[i].text:
            docx.paragraphs[i].runs[1].text = MJnum
        if u'时    间' in docx.paragraphs[i].text:
            docx.paragraphs[i].runs[1].text = doctime


###项目信息收集，基本是第一个表格
project_info = []
for i in range(len(docx1.tables[0].rows)):
    for j in range(len(docx1.tables[0].rows[i].cells)):
        project_info.append(docx1.tables[0].rows[i].cells[j].text)

project_info = '\t'.join(project_info)
project_name = re.findall(u'项目名称\t项目名称\t项目名称\t项目名称\t(.*?)\t', project_info)[0]
contract_no = re.findall(u'合同编号\t合同编号\t合同编号\t合同编号\t(.*?)\t', project_info)[0]
apecies_information = re.findall(u'物种信息\t(.*?)\t', project_info)[0]
exp_purpose = re.findall(u'实验目的\t(.*?)\t', project_info)[0]
client_name = re.findall(u'单位名称\t(.*?)\t', project_info)[0]
client_address = re.findall(u'单位地址\t(.*?)\t', project_info)[0]
tutor_info = re.findall(u'实验室导师\t(.*?)\t电话\t(.*?)\t实验室导师\t(.*?)\t邮箱\t(.*?)\t',
                        project_info)[0]
contactor_info = re.findall(
    u'项目联系人\t(.*?)\t电话\t(.*?)\t项目联系人\t(.*?)\t邮箱\t(.*?)\t', project_info)[0]
seller_info = re.findall(u'销售员\t(.*?)\t电话\t(.*?)\t销售员\t(.*?)\t邮箱\t(.*?)\t',
                         project_info)[0]
supporter = re.findall(u'技术支持\t(.*?)\t电话\t(.*?)\t技术支持\t(.*?)\t邮箱\t(.*)',
                       project_info)[0]


###项目信息登记，好像只能一个一个写入
def _project_info(docx):
    for i in range(len(docx.tables)):
        if u'项目名称' in docx.tables[i].rows[0].cells[0].text:
            font_size = docx.tables[i].rows[0].cells[0].paragraphs[0].runs[
                0].font.size
            font_name = docx.tables[i].rows[0].cells[0].paragraphs[0].runs[
                0].font.name
            for j in range(len(docx.tables[i].rows)):
                if u'项目名称' in docx.tables[i].rows[j].cells[0].text:
                    run = docx.tables[i].rows[j +
                                              1].cells[0].paragraphs[0].runs[0]
                    run.text = project_name
                    run.font.size = font_size
                    run.font.name = font_name
                if u'合同编号' in docx.tables[i].rows[j].cells[0].text:
                    run = docx.tables[i].rows[j +
                                              1].cells[0].paragraphs[0].runs[0]
                    run.text = contract_no
                    run.font.size = font_size
                    run.font.name = font_name
                if u'物种信息' in docx.tables[i].rows[j].cells[0].text:
                    run = docx.tables[i].rows[j].cells[1].paragraphs[0].runs[0]
                    run.text = apecies_information
                    run.font.size = font_size
                    run.font.name = font_name
                if u'实验目的' in docx.tables[i].rows[j].cells[0].text:
                    run = docx.tables[i].rows[j].cells[1].paragraphs[0].runs[0]
                    run.text = exp_purpose
                    run.font.size = font_size
                    run.font.name = font_name
                if u'单位名称' in docx.tables[i].rows[j].cells[0].text:
                    run = docx.tables[i].rows[j].cells[1].paragraphs[0].runs[0]
                    run.text = exp_purpose
                    run.font.size = font_size
                    run.font.name = font_name
                if u'单位地址' in docx.tables[i].rows[j].cells[0].text:
                    run = docx.tables[i].rows[j].cells[1].paragraphs[0].runs[0]
                    run.text = exp_purpose
                    run.font.size = font_size
                    run.font.name = font_name
                if u'实验室导师' in docx.tables[i].rows[j].cells[0].text:
                    if u'电话' in docx.tables[i].rows[j].cells[2].text:
                        run_teacher = docx.tables[i].rows[j].cells[
                            1].paragraphs[0].runs[0]
                        run_teacher.text = tutor_info[0]
                        run_teacher.font.size = font_size
                        run_teacher.font.name = font_name
                        run_phone = docx.tables[i].rows[j].cells[3].paragraphs[
                            0].runs[0]
                        run_phone.text = tutor_info[1]
                        run_phone.font.size = font_size
                        run_phone.font.name = font_name
                        run_email = docx.tables[i].rows[j + 1].cells[
                            3].paragraphs[0].runs[0]
                        run_email.text = tutor_info[3]
                        run_email.font.size = font_size
                        run_email.font.name = font_name
                if u'项目联系人' in docx.tables[i].rows[j].cells[0].text:
                    if u'电话' in docx.tables[i].rows[j].cells[2].text:
                        run_contactor = docx.tables[i].rows[j].cells[
                            1].paragraphs[0].runs[0]
                        run_contactor.text = contactor_info[0]
                        run_contactor.font.size = font_size
                        run_contactor.font.name = font_name
                        run_phone = docx.tables[i].rows[j].cells[3].paragraphs[
                            0].runs[0]
                        run_phone.text = contactor_info[1]
                        run_phone.font.size = font_size
                        run_phone.font.name = font_name
                        run_email = docx.tables[i].rows[j + 1].cells[
                            3].paragraphs[0].runs[0]
                        run_email.text = contactor_info[3]
                        run_email.font.size = font_size
                        run_email.font.name = font_name
                if u'销售员' in docx.tables[i].rows[j].cells[0].text:
                    if u'电话' in docx.tables[i].rows[j].cells[2].text:
                        run_seller = docx.tables[i].rows[j].cells[
                            1].paragraphs[0].runs[0]
                        run_seller.text = seller_info[0]
                        run_seller.font.size = font_size
                        run_seller.font.name = font_name
                        run_phone = docx.tables[i].rows[j].cells[3].paragraphs[
                            0].runs[0]
                        run_phone.text = seller_info[1]
                        run_phone.font.size = font_size
                        run_phone.font.name = font_name
                        run_email = docx.tables[i].rows[j + 1].cells[
                            3].paragraphs[0].runs[0]
                        run_email.text = seller_info[3]
                        run_email.font.size = font_size
                        run_email.font.name = font_name
                if u'技术支持' in docx.tables[i].rows[j].cells[0].text:
                    if u'电话' in docx.tables[i].rows[j].cells[2].text:
                        run_supporter = docx.tables[i].rows[j].cells[
                            1].paragraphs[0].runs[0]
                        run_supporter.text = supporter[0]
                        run_supporter.font.size = font_size
                        run_supporter.font.name = font_name
                        run_phone = docx.tables[i].rows[j].cells[3].paragraphs[
                            0].runs[0]
                        run_phone.text = supporter[1]
                        run_phone.font.size = font_size
                        run_phone.font.name = font_name
                        run_email = docx.tables[i].rows[j + 1].cells[
                            3].paragraphs[0].runs[0]
                        run_email.text = supporter[3]
                        run_email.font.size = font_size
                        run_email.font.name = font_name


###对表格进行写入:__MJ_table执行思路是判断第一行的第一列和第二列的信息来判断表格属于谁
def __MJ_table(i, docx, header1, header2, file, judgment1, judgment2,
               table_format, lines):
    file = commands.getoutput('''ls %s |head -1 ''' % file)
    if u'%s' % header1 in docx.tables[i].rows[0].cells[
            0].text and u'%s' % header2 in docx.tables[i].rows[0].cells[
                1].text:
        with open(file, 'r') as table:
            mark = 1
            for line in table.readlines():
                if mark <= lines:
                    items = line.strip().split('\t')
                    count = re.subn(r'%s' % judgment1, r'%s' % judgment1,
                                    line)[1]
                    if 0 < count and count < judgment2:
                        for j in range(len(docx.tables[i].rows[mark].cells)):
                            run = docx.tables[i].rows[mark].cells[
                                j].paragraphs[0].runs[0]
                            font_name = run.font.name
                            font_size = run.font.size
                            run.text = items[j]
                            run.font.size = font_size
                            run.font.name = font_name
                            docx.tables[i].rows[mark].cells[j].paragraphs[
                                0].paragraph_format.alignment = table_format
                        mark += 1


def _MJ_tables(docx):
    for i in range(len(docx.tables)):
        table_format = docx.tables[i].rows[0].cells[0].paragraphs[
            0].paragraph_format.alignment
        #4.1.1
        ##GO.list
        __MJ_table(i, docx, u'蛋白Accession号', u'对应的GO编号',
                   '%s/*Annotation/*GO/GO.list' % (args.file), r'GO:', 5,
                   table_format, 2)
        ##*level2/3/4.xls
        __MJ_table(i, docx, u'GO注释分类的分支', u'GO分类的定义',
                   '%s/*Annotation/*GO/GO.list.level2.xls' % (args.file),
                   r'GO:', 10, table_format, 2)
        #4.1.2
        ##pathway.txt
        __MJ_table(i, docx, u'蛋白的Accession编号', u'对应的KO号',
                   '%s/*Annotation/*KEGG/pathway.txt' % (args.file), r'K', 10,
                   table_format, 2)
        ##pathway_table.xls
        __MJ_table(i, docx, u'通路的编号', u'通路的定义',
                   '%s/*Annotation/*KEGG/pathways/pathway_table.xls' %
                   (args.file), r'K', 5, table_format, 2)
        ##kegg_table.xls
        __MJ_table(i, docx, u'蛋白名称', u'KO编号',
                   '%s/*Annotation/*KEGG/pathways/kegg_table.xls' %
                   (args.file), r'ko', 5, table_format, 2)
        #4.1.3-----COG
        ##COG.list
        __MJ_table(i, docx, u'蛋白Accession号', u'对应的COG号',
                   '%s/*Annotation/*COG/COG.list' % (args.file), r'COG\d+', 5,
                   table_format, 2)
        ##COG.annot.xls
        __MJ_table(i, docx, u'蛋白名称', u'COG编号',
                   '%s/*Annotation/*COG/COG.annot.xls' % (args.file), r'\[',
                   10, table_format, 2)
        ##COG.class.catalog.xls
        __MJ_table(i, docx, u'COG 4种类型', u'COG功能分类，共25',
                   '%s/*Annotation/*COG/COG.class.catalog.xls' % (args.file),
                   r'\[', 10, table_format, 2)
        #4.1.3-----KOG
        ##KOG.list
        __MJ_table(i, docx, u'蛋白Accession号', u'对应的KOG号',
                   '%s/*Annotation/*KOG/KOG.list' % (args.file), r'KOG\d+', 5,
                   table_format, 2)
        ##COG.annot.xls
        __MJ_table(i, docx, u'蛋白名称', u'KOG编号',
                   '%s/*Annotation/*KOG/KOG.annot.xls' % (args.file), r'\[',
                   10, table_format, 2)
        ##COG.class.catalog.xls
        __MJ_table(i, docx, u'KOG 4种类型', u'KOG功能分类，共25',
                   '%s/*Annotation/*KOG/KOG.class.catalog.xls' % (args.file),
                   r'\[', 10, table_format, 2)
        #4.2.1
        ##all_diff_up_down.xls
        __MJ_table(i, docx, u'名称', u'蛋白总数',
                   '%s/*DiffExpAnalysis/*Statistics/all_diff_up_down.xls' %
                   (args.file), r'vs', 10, table_format, 6)
        ##*_vs_*.diff.exp.xls
        __MJ_table(
            i, docx, u'蛋白Accession编号', u'样本1中该蛋白相对表达量均值',
            '%s/*DiffExpAnalysis/*Statistics/Volcano/*_vs_*.diff.exp.xls' %
            (args.file), r'\.', 5, table_format, 2)
        #4.2.2
        ##*.enrichment.detail.xls
        __MJ_table(i, docx, u'Id', u'Enrichment',
                   '%s/*DiffExpAnalysis/*GO/Enrichment/*_vs_*.detail.xls' %
                   (args.file), r';', 4, table_format, 2)
        #4.2.3
        ##*.pathway.xls
        __MJ_table(i, docx, u'KEGG pathway名称', u'数据库',
                   '%s/*DiffExpAnalysis/*KEGG/Enrichment/*.pathway.xls' %
                   (args.file), r'ko', 5, table_format, 2)
        #4.2.5
        ##*_vs_*.network.xls
        __MJ_table(i, docx, r'node1', r'node2',
                   '%s/*DiffExpAnalysis/*Network/*interaction.xls' %
                   (args.file), r'\.', 5, table_format, 2)
        ##*_vs_*.annotation.xls
        __MJ_table(i, docx, u'蛋白名', u'蛋白uniprot登录号',
                   '%s/*DiffExpAnalysis/*Network/*annotation.xls' %
                   (args.file), r'\.', 5, table_format, 2)


###对图片进行替换，由图片上一句话来进行判断图片的归属信息
def __MJ_jpg(i, docx, mark, file, density, x, y, dx, dy, width):
    if u'%s' % mark in docx.paragraphs[i].text:
        file_ls = commands.getoutput('''/bin/ls %s |head -1 ''' % file)
        file_ls = file_ls.strip('\n')
        if not r'cannot access' in file_ls:
            file_pdf = file_ls
            file_jpg = 'tmp/%s.png' % (file_ls.split('/')[-1])
            if 'pdf' in file_ls:
                os.system('''convert -density %s %s %s >/dev/null 2>&1''' %
                          (density, file_pdf, file_jpg))
            elif 'network' in mark or r'Ipath' in mark:
                img = Image.open(file_pdf)
                y = img.size[1] / (img.size[0] / x)
                os.system('''convert -resize %sx%s %s %s >/dev/null 2>&1''' %
                          (x, y, file_pdf, file_jpg))
            else:
                file_jpg = file_pdf
            if 'subclusters' in mark:
                file_jpg = commands.getoutput(
                    '''ls tmp/*subclusters*.png |head -1 ''').strip()
            if u'KEGG通路图片' in mark or 'network' in mark or r'Ipath' in mark:
                print u'\t%s\n\t\t图片没有裁剪' % mark
            else:
                os.system('''convert %s -crop %sx%s+%s+%s %s''' %
                          (file_jpg, x, y, dx, dy, file_jpg))
            docx.paragraphs[i + 1].add_run()
            docx.paragraphs[i + 1].runs[1].style.style_id = docx.paragraphs[
                i + 1].runs[0].style.style_id
            docx.paragraphs[i + 1].runs[1].add_picture(file_jpg, width=width)
            docx.paragraphs[i + 1].runs[0].clear()
        else:
            print u'\t%s\n\t\t没有这张图片' % mark


def _MJ_jpgs(docx):
    if os.path.isdir('tmp'):
        os.system('/bin/rm -rf tmp')
    os.mkdir('tmp')
    for i in range(len(docx.paragraphs)):
        #QC
        ##肽段匹配误差分布图
        __MJ_jpg(i, docx, u'dMass.pdf：肽段匹配误差分布图',
                 '%s/*QualityControl/dMass.pdf' % (args.file), 150, 1799, 1199,
                 0, 0, 4800000)
        ##肽段数量分布柱状图
        __MJ_jpg(i, docx, u'Peptide_number_distribution.pdf：肽段数量分布柱状图',
                 '%s/*QualityControl/Peptide_number_distribution.pdf' %
                 (args.file), 150, 1499, 1049, 0, 0, 4800000)
        ##肽段长度分布柱状图
        __MJ_jpg(i, docx, u'Peptide_length_distribution.pdf：肽段长度分布柱状图',
                 '%s/*QualityControl/Peptide_length_distribution.pdf' %
                 (args.file), 150, 1499, 1049, 0, 0, 4800000)
        ##蛋白分子量分布柱状图
        __MJ_jpg(
            i, docx, u'Protein_molecular_weight_distribution.pdf：蛋白分子量分布柱状图',
            '%s/*QualityControl/Protein_molecular_weight_distribution.pdf' %
            (args.file), 150, 1499, 1049, 0, 0, 4800000)
        ##蛋白覆盖度分布饼图
        __MJ_jpg(i, docx, u'Protein_coverage_distribution.pdf：蛋白覆盖度分布饼图',
                 '%s/*QualityControl/Protein_coverage_distribution.pdf' %
                 (args.file), 150, 1119, 919, 350, 120, 3800000)
        ##鉴定蛋白质信息统计柱状图
        __MJ_jpg(i, docx, u'Protein_information.pdf：鉴定蛋白质信息统计柱状图',
                 '%s/*QualityControl/Protein_information.pdf' % (args.file),
                 150, 1499, 1049, 0, 0, 4800000)
        #4.1.1
        ##GO二级分类统计条形图
        __MJ_jpg(i, docx, u'level2.go.txt.pdf：GO二级分类统计条形图',
                 '%s/*Annotation/*GO/level2.go.txt.pdf' % (args.file), 150,
                 1500, 1100, 80, 150, 4800000)
        ##GO二级、三级、四级分类统计九饼图
        __MJ_jpg(i, docx, u'level234.pdf：GO二级、三级、四级分类统计九饼图',
                 '%s/*Annotation/*GO/GO.list.Level234.pdf' % (args.file), 75,
                 1500, 900, 0, 150, 4800000)
        #4.1.2
        ##包含蛋白数目最多的前20个通路
        __MJ_jpg(i, docx, u'pathway.top20.pdf：包含蛋白数目最多的前20个通路（除以ko01开头的基础通路外）',
                 '%s/*Annotation/*KEGG/pathway.top20.pdf' % (args.file), 150,
                 1799, 1199, 0, 0, 4800000)
        ##KEGG通路图片展示
        __MJ_jpg(i, docx, u'pathways文件夹下的.png文件：KEGG通路图片展示',
                 '%s/*Annotation/*KEGG/pathways/ko*[345]*.png' % (args.file),
                 150, 0, 0, 0, 0, 4000000)
        ##KEGG代谢通路进行分类
        __MJ_jpg(i, docx,
                 u'kegg_classification.pdf：对蛋白做KO注释后，可根据它们参与的KEGG代谢通路进行分类',
                 '%s/*Annotation/*KEGG/kegg_classification.pdf' % (args.file),
                 150, 1200, 900, 1, 400, 4000000)
        #4.1.3
        ##COG功能分类统计柱图
        __MJ_jpg(i, docx, u'COG.class.catalog.pdf：COG功能分类统计柱图',
                 '%s/*Annotation/*COG/COG.class.catalog.pdf' % (args.file),
                 150, 1499, 850, 0, 1, 5200000)
        ##KOG功能分类统计图
        __MJ_jpg(i, docx, u'KOG.class.catalog.pdf：KOG功能分类统计柱图',
                 '%s/*Annotation/*KOG/KOG.class.catalog.pdf' % (args.file),
                 150, 1499, 850, 0, 0, 5200000)
        #4.2.1
        ##差异蛋白可视化火山图
        __MJ_jpg(i, docx, u'各分组样本差异蛋白可视化火山图',
                 '%s/*DiffExpAnalysis/*Statistics/Volcano/*_vs_*.volcano.pdf' %
                 (args.file), 150, 1150, 820, 90, 40, 4800000)
        ##差异蛋白可视化散点图
        __MJ_jpg(i, docx, u'各分组样本差异蛋白可视化散点图',
                 '%s/*DiffExpAnalysis/*Statistics/Volcano/*_vs_*.scatter.pdf' %
                 (args.file), 150, 1150, 820, 90, 40, 4800000)
        ##差异蛋白Venn图
        __MJ_jpg(i, docx, u'*.Venn.pdf：差异蛋白Venn图',
                 '%s/*DiffExpAnalysis/*Statistics/Venn/*Venn.pdf' %
                 (args.file), 150, 1250, 1250, 150, 150, 4000000)
        #4.2.2
        ##上下调蛋白GO注释柱形图
        __MJ_jpg(i, docx, u'*gobars.pdf：上下调蛋白GO注释柱形图',
                 '%s/*DiffExpAnalysis/*GO/Annotation/*gobars.pdf' %
                 (args.file), 150, 1700, 1500, 50, 0, 4800000)
        ##GO功能富集分析柱状图
        __MJ_jpg(i, docx, u'*.enrichment.detail.xls.go.pdf：各分组差异蛋白GO功能富集分析柱状图',
                 '%s/*DiffExpAnalysis/*GO/Enrichment/*.go.pdf' % (args.file),
                 150, 0, 0, 0, 0, 6300000)
        #4.2.3
        ##差异蛋白KEGG通路图片展示
        __MJ_jpg(
            i, docx, u'*.png：差异蛋白KEGG通路图片展示',
            '%s/*DiffExpAnalysis/*KEGG/Annotation/*_vs_*.diff.exp.xls.path/ko*[345]*.png'
            % (args.file), 150, 0, 0, 0, 0, 4000000)
        ##分组差异蛋白KEGG pathway富集分析柱状图
        __MJ_jpg(i, docx, u'*.pathway.pdf：各分组差异蛋白KEGG pathway富集分析柱状图',
                 '%s/*DiffExpAnalysis/*KEGG/Enrichment/*pathway.pdf' %
                 (args.file), 150, 0, 0, 0, 0, 6300000)
        #4.2.4
        ##差异蛋白表达模式聚类图
        __MJ_jpg(i, docx, u'Heatmap.pdf：差异蛋白表达模式聚类图',
                 '%s/*DiffExpAnalysis/*Cluster/Heatmap.pdf' % (args.file), 150,
                 1799, 900, 0, 0, 4800000)
        ##差异蛋白模块（clusters）表达趋势折线图
        __MJ_jpg(
            i, docx,
            u'Heatmap_trendlines_for_*_subclusters.pdf：差异蛋白模块（clusters）表达趋势折线图',
            '%s/*DiffExpAnalysis/*Cluster/*subclusters.pdf' % (args.file), 150,
            1049, 1049, 0, 0, 3800000)
        #4.2.5
        ##两组样本之间差异蛋白的蛋白互作图（png格式）
        __MJ_jpg(i, docx, u'*_vs_*.network.png：两组样本之间差异蛋白的蛋白互作图（png格式）',
                 '%s/*DiffExpAnalysis/*Network/*confidence.png' % (args.file),
                 0, 1200, 0, 0, 0, 5200000)
        #4.2.6
        ##Ipath整合通路图
        __MJ_jpg(i, docx, u'*_vs_*.Ipath.png：Ipath整合通路图（png格式）',
                 '%s/*DiffExpAnalysis/*Ipath/*ipath.png' % (args.file), 0, 800,
                 0, 0, 0, 4800000)


def _MJ_up(docx):
    up = float(args.up)
    up = '%.2f' % (up)
    down = '%.2f' % (1 / float(up))
    for i in range(len(docx.paragraphs)):
        if u'显著差异表达蛋白的筛选标准为' in docx.paragraphs[i].text:
            docx.paragraphs[i].add_run()
            run0 = docx.paragraphs[i].runs[0]
            run = docx.paragraphs[i].runs[1]
            run.text = u'（FC＜%s或 FC＞%s）。' % (down, up)
            run.font.size = run0.font.size
            run.font.name = run0.font.name


print 'M2.docx是质控:'
_first_page(docx2)
_project_info(docx2)
_MJ_jpgs(docx2)
docx2.save('M2.docx')
print '\tM2.docx is done!\nM3.docx是生信分析:'
_first_page(docx3)
_project_info(docx3)
_MJ_up(docx3)
_MJ_tables(docx3)
_MJ_jpgs(docx3)
docx3.save('M3.docx')
print '\tM2.docx is done!\nall done!'
