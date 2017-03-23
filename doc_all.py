#!/usr/bin/env python
#encoding:utf-8
#jun.yan@majorbio.com
#20170320

import argparse
import commands
import os
import sys
import re
from docx import Document

parser = argparse.ArgumentParser(description = "iTRAQ prok report automatically generate scripts")
parser.add_argument("--org",dest = "org",type = str,required=True,help ="'ork' or 'prok'")
parser.add_argument("-m","--test_department_m1",dest = "M1",type = str,required=True,help ="test_department report file")
parser.add_argument("-f","--file",dest = "file",type = str,required = True,help = "please input file,normally 'Result_Files'")
args = parser.parse_args()

docx1 = Document(args.M1)
docx2 = Document('/mnt/ilustre/users/jun.yan/scripts/pm_report_module/iTRAQ/%s/iTRAQ2.docx' %args.org)
docx3 = Document('/mnt/ilustre/users/jun.yan/scripts/pm_report_module/iTRAQ/%s/iTRAQ3.docx' %args.org)

text1 = '\n'.join([docx1.paragraphs[i].text for i in range(len(docx1.paragraphs))])

def _first_page(docx):
	name = re.findall(u'�ͻ�������(.*)',text1)[0]
	print name
	MJnum = re.findall(u'��Ŀ��ţ�(.*)',text1)[0]
	print MJnum
	doctime = re.findall(u'ʱ    �䣺(.*)',text1)[0]
	print doctime
	for i in range(len(docx.paragraphs)):
		if u'�ͻ�����' in docx.paragraphs[i].text:
			print len(docx.paragraphs[i].runs)
			docx.paragraphs[i].runs[1].text = name
		if u'��Ŀ���' in docx.paragraphs[i].text:
			docx.paragraphs[i].runs[1].text = MJnum
		if u'ʱ    ��' in docx.paragraphs[i].text:
			docx.paragraphs[i].runs[1].text = doctime

_first_page(docx2)
_first_page(docx3)

def _project_info(docx):
	project_info = []
	for i in range(len(docx1.tables[0].rows)):
		for j in range(len(docx1.tables[0].rows[i].cells)):
			project_info.append(docx1.tables[0].rows[i].cells[j].text)
	project_info = '\t'.join(project_info)
	
	project_name = re.findall(u'��Ŀ����\t��Ŀ����\t��Ŀ����\t��Ŀ����\t(.*?)\t',project_info)[0]
	print project_name
	contract_no = re.findall(u'��ͬ���\t��ͬ���\t��ͬ���\t��ͬ���\t(.*?)\t',project_info)[0]
	print contract_no
	apecies_information = re.findall(u'������Ϣ\t(.*?)\t',project_info)[0]
	print apecies_information
	exp_purpose = re.findall(u'ʵ��Ŀ��\t(.*?)\t',project_info)[0]
	print exp_purpose
	client_name = re.findall(u'��λ����\t(.*?)\t',project_info)[0]
	print client_name
	client_address = re.findall(u'��λ��ַ\t(.*?)\t',project_info)[0]
	print client_address
	tutor_info =  re.findall(u'ʵ���ҵ�ʦ\t(.*?)\t�绰\t(.*?)\tʵ���ҵ�ʦ\t(.*?)\t����\t(.*?)\t',project_info)[0]
	print tutor_info
	contactor_info = re.findall(u'��Ŀ��ϵ��\t(.*?)\t�绰\t(.*?)\t��Ŀ��ϵ��\t(.*?)\t����\t(.*?)\t',project_info)[0]
	print contactor_info
	seller_info =re.findall(u'����Ա\t(.*?)\t�绰\t(.*?)\t����Ա\t(.*?)\t����\t(.*?)\t',project_info)[0]
	print seller_info
	supporter = re.findall(u'����֧��\t(.*?)\t�绰\t(.*?)\t����֧��\t(.*?)\t����\t(.*)',project_info)[0]
	print supporter
	
	for i in range(len(docx.tables)):
		for j in range(len(docx.tables[i].rows)):
			if u'��Ŀ����' in docx.tables[i].rows[j].cells[0].text:
				docx.tables[i].rows[j+1].cells[0].text = project_name
			if u'��ͬ���' in docx.tables[i].rows[j].cells[0].text:
				docx.tables[i].rows[j+1].cells[0].text = contract_no
			if u'������Ϣ' in docx.tables[i].rows[j].cells[0].text:
				docx.tables[i].rows[j].cells[1].text = apecies_information
			if u'ʵ��Ŀ��' in docx.tables[i].rows[j].cells[0].text:
				docx.tables[i].rows[j].cells[1].text = exp_purpose
			if u'��λ����' in docx.tables[i].rows[j].cells[0].text:
				docx.tables[i].rows[j].cells[1].text = client_name
			if u'��λ��ַ' in docx.tables[i].rows[j].cells[0].text:
				docx.tables[i].rows[j].cells[1].text = client_address
			if u'ʵ���ҵ�ʦ' in docx.tables[i].rows[j].cells[0].text:
				if u'�绰' in docx.tables[i].rows[j].cells[2].text :
					docx.tables[i].rows[j].cells[1].text = tutor_info[0]
					docx.tables[i].rows[j].cells[3].text = tutor_info[1]
					docx.tables[i].rows[j+1].cells[3].text = tutor_info[3]
			if u'��Ŀ��ϵ��' in docx.tables[i].rows[j].cells[0].text:	
				if u'�绰' in docx.tables[i].rows[j].cells[2].text :			
					docx.tables[i].rows[j].cells[1].text = contactor_info[0]
					docx.tables[i].rows[j].cells[3].text = contactor_info[1]
					docx.tables[i].rows[j+1].cells[3].text = contactor_info[3]
			if u'����Ա' in docx.tables[i].rows[j].cells[0].text:
				if u'�绰' in docx.tables[i].rows[j].cells[2].text :
					docx.tables[i].rows[j].cells[1].text = seller_info[0]
					docx.tables[i].rows[j].cells[3].text = seller_info[1]
					docx.tables[i].rows[j+1].cells[3].text = seller_info[3]
			if u'����֧��' in docx.tables[i].rows[j].cells[0].text:
				if u'�绰' in  docx.tables[i].rows[j].cells[2].text :
					docx.tables[i].rows[j].cells[1].text = supporter[0]
					docx.tables[i].rows[j].cells[3].text = supporter[1]
					docx.tables[i].rows[j+1].cells[3].text = supporter[3]

_project_info(docx2)
_project_info(docx3)

docx2 = Document('/mnt/ilustre/users/jun.yan/scripts/pm_report_module/iTRAQ/prok/iTRAQ2.docx')		
docx3 = Document('/mnt/ilustre/users/jun.yan/scripts/pm_report_module/iTRAQ/prok/iTRAQ3.docx')	
				
def __MJ_table(i,docx,header1,header2,file,judgment1,judgment2,table_format): 
	file = commands.getoutput('''ls %s |head -1 ''' %file)
	if u'%s' %header1 in docx.tables[i].rows[0].cells[0].text and u'%s' %header2 in docx.tables[i].rows[0].cells[1].text:
		with open(file,'r') as table:
			mark = 1
			for line in table.readlines():
				if mark<3:
					items = line.strip().split('\t')
					count = re.subn(r'%s' %judgment1,r'%s' %judgment1,line)[1]
					if count<judgment2:
						for j in range(len(docx.tables[i].rows[mark].cells)):
							docx.tables[i].rows[mark].cells[j].paragraphs[0].add_run()
							run=docx.tables[i].rows[mark].cells[j].paragraphs[0].runs[1]
							run.text=items[j]
							run.font.size = docx.tables[i].rows[mark].cells[j].paragraphs[0].runs[0].font.size
							run.font.name = docx.tables[i].rows[mark].cells[j].paragraphs[0].runs[0].font.name
							docx.tables[i].rows[mark].cells[j].paragraphs[0].runs[0].clear()
							docx.tables[i].rows[mark].cells[j].paragraphs[0].paragraph_format.alignment=table_format
						mark += 1

def _MJ3_tables(docx):
	for i in range(len(docx.tables)):
		table_format = docx.tables[i].rows[0].cells[0].paragraphs[0].paragraph_format.alignment
		#4.1.1
		##GO.list
		__MJ_table(i,docx,u'����Accession��',u'��Ӧ��GO���','%s/2.Annotation/2.1.GO/GO.list',r'GO:' %args.file,5,table_format)
		##*level2/3/4.xls
		__MJ_table(i,docx,u'GOע�ͷ���ķ�֧',u'GO����Ķ���','%s/2.Annotation/2.1.GO/GO.list.level2.xls'%args.file,r'GO:',10,table_format)
		#4.1.2
		##pathway.txt
		__MJ_table(i,docx,u'���׵�Accession���',u'��Ӧ��KO��','%s/2.Annotation/2.2.KEGG/pathway.txt'%args.file,r'KO:',10,table_format)
		##pathway_table.xls
		__MJ_table(i,docx,u'ͨ·�ı��',u'ͨ·�Ķ���','%s/2.Annotation/2.2.KEGG/pathways/pathway_table.xls'%args.file,r'KO:',5,table_format)
		##kegg_table.xls
		__MJ_table(i,docx,u'��������',u'KO���','%s/2.Annotation/2.2.KEGG/pathways/kegg_table.xls'%args.file,r'ko:',5,table_format)
		#4.1.3
		##COG.list
		__MJ_table(i,docx,u'����Accession��',u'��Ӧ��COG��','%s/2.Annotation/2.3.COG/COG.list'%args.file,r'COG:',5,table_format)
		##COG.annot.xls
		__MJ_table(i,docx,u'��������',u'COG���','%s/2.Annotation/2.3.COG/COG.annot.xls'%args.file,r'COG:',10,table_format)
		##COG.class.catalog.xls
		__MJ_table(i,docx,u'COG 4������',u'COG���ܷ��࣬��25','%s/2.Annotation/2.3.COG/COG.class.catalog.xls'%args.file,r'\[',10,table_format)
		##*_vs_*.diff.exp.xls
		__MJ_table(i,docx,u'����Accession���',u'����1�иõ�����Ա������ֵ','%s/3.DiffExpAnalysis/3.1.Statistics/Volcano/*_vs_*.diff.exp.xls'%args.file,r'\.',5,table_format)
		#4.2.3
		##*.enrichment.detail.xls
		__MJ_table(i,docx,u'Id',u'Enrichment','%s/3.DiffExpAnalysis/3.2.GO/Enrichment/*_vs_*.diff.exp.xls.DE.list.enrichment.detail.xls'%args.file,r'GO:',5,table_format)
		##*.pathway.xls
		__MJ_table(i,docx,u'KEGG pathway����',u'���ݿ�','%s/3.DiffExpAnalysis/3.3.KEGG/Enrichment/*.pathway.xls'%args.file,r'ko:',5,table_format)

_MJ3_tables(docx3)

def __MJ_jpg(i,docx,mark,file):
	file = commands.getoutput('''ls %s |head -1 ''' %file)
	if u'%s' %mark in docx.paragraphs[i]:	
		if 'pdf' in str(file):
			file_pdf = file
			file_jpg = 'tmp/tmp_file.jpg'
			os.system('''convert -density 300  %s %s >/dev/null 2>&1''' %(file_pdf,file_jpg))
		else:
			file_jpg = file
		docx2.paragraphs[i+1].add_run()
		docx2.paragraphs[i+1].runs[1].style.style_id = docx2.paragraphs[i+1].runs[0].style.style_id
		docx2.paragraphs[i+1].runs[1].add_picture(file_jpg,width = 4444444)
		docx2.paragraphs[i+1].runs[0].clear()

def _MJ3_jpgs(docx):
	if not os.path.isdir('tmp'):
		os.mkdir('tmp')
	for i in range(len(docx.paragraphs)):
		#4.1.1
		##GO��������ͳ������ͼ
		__MJ_jpg(i,docx,u'level2.go.txt.pdf��GO��������ͳ������ͼ������ͼ','Result_Files/2.Annotation/2.1.GO/level2.go.txt.pdf')
		##GO�������������ļ�����ͳ�ƾű�ͼ
		__MJ_jpg(i,docx,u'level234.pdf��GO�������������ļ�����ͳ�ƾű�ͼ������ͼ','Result_Files/2.Annotation/2.1.GO/GO.list.Level234.pdf')
		#4.1.2
		##����������Ŀ����ǰ20��ͨ·
		__MJ_jpg(i,docx,u'pathway.top20.pdf������������Ŀ����ǰ20��ͨ·������ͼ','Result_Files/2.Annotation/2.2.KEGG/pathway.top20.pdf')
		##KEGGͨ·ͼƬչʾ
		__MJ_jpg(i,docx,u'pathways�ļ����µ�.png�ļ���KEGGͨ·ͼƬչʾ������ͼ','Result_Files/2.Annotation/2.2.KEGG/pathways/*.png')
		##KEGG��лͨ·���з���
		__MJ_jpg(i,docx,u'kegg_classification.pdf���Ե�����KOע�ͺ󣬿ɸ������ǲ����KEGG��лͨ·���з���','Result_Files/2.Annotation/2.2.KEGG/kegg_classification.pdf')
		#4.1.3
		##COG���ܷ���ͳ����ͼ
		__MJ_jpg(i,docx,u'COG.class.catalog.pdf��COG���ܷ���ͳ����ͼ������ͼ','Result_Files/2.Annotation/2.3.COG/COG.class.catalog.pdf')		
		#4.2.1
		##���쵰�׿��ӻ���ɽͼ
		__MJ_jpg(i,docx,u'*_vs_*.volcano.pdf���������������쵰�׿��ӻ���ɽͼ������ͼ','Result_Files/3.DiffExpAnalysis/3.1.Statistics/Volcano/*_vs_*.volcano.pdf')
		##���쵰��Vennͼ
		__MJ_jpg(i,docx,u'*.Venn.pdf�����쵰��Vennͼ������ͼ','Result_Files/3.DiffExpAnalysis/3.1.Statistics/Venn/*Venn.pdf')
		##���µ�����GOע������ͼ		
		__MJ_jpg(i,docx,u'*gobars.pdf�����µ�����GOע������ͼ������ͼ','Result_Files/3.DiffExpAnalysis/3.2.GO/Annotation/*.listlevel2-gobars.pdf')
		#4.2.2
		##GO���ܸ���������״ͼ
		__MJ_jpg(i,docx,u'*.enrichment.detail.xls.go.pdf����������쵰��GO���ܸ���������״ͼ������ͼ','Result_Files/3.DiffExpAnalysis/3.2.GO/Enrichment/*.go.pdf')
		#4.2.3
		##���쵰��KEGGͨ·ͼƬչʾ
		__MJ_jpg(i,docx,u'*.png�����쵰��KEGGͨ·ͼƬչʾ������ͼ','Result_Files/3.DiffExpAnalysis/3.3.KEGG/Annotation/*_vs_*.diff.exp.xls.path/*.png')
		##������쵰��KEGG pathway����������״ͼ
		__MJ_jpg(i,docx,u'*.pathway.pdf����������쵰��KEGG pathway����������״ͼ������ͼ','Result_Files/3.DiffExpAnalysis/3.3.KEGG/Enrichment/*pathway.pdf')
		#4.2.4
		##���쵰�ױ��ģʽ����ͼ
		__MJ_jpg(i,docx,u'Heatmap.pdf�����쵰�ױ��ģʽ����ͼ','Result_Files/3.DiffExpAnalysis/3.4.Cluster/Heatmap.pdf')
		##���쵰��ģ�飨clusters�������������ͼ
		__MJ_jpg(i,docx,u'Heatmap_trendlines_for_*_subclusters.pdf�����쵰��ģ�飨clusters�������������ͼ','Result_Files/3.DiffExpAnalysis/3.4.Cluster/Heatmap_trendlines_for_10_subclusters.pdf')
		#4.2.5
		##��������֮����쵰�׵ĵ��׻���ͼ��png��ʽ��
		__MJ_jpg(i,docx,u'*_vs_*.network.png����������֮����쵰�׵ĵ��׻���ͼ��png��ʽ��������ͼ','Result_Files/3.DiffExpAnalysis/3.5.Network/*confidence.png')
		#4.2.6
		##Ipath����ͨ·ͼ
		__MJ_jpg(i,docx,u'*_vs_*.Ipath.png��Ipath����ͨ·ͼ��png��ʽ��������ͼ','Result_Files/3.DiffExpAnalysis/3.6.Ipath/*ipath.png')
		
	os.system('rm -rf tmp')	

_MJ3_jpgs(docx3)


docx2.save('M2.docx')
docx3.save('M3.docx')